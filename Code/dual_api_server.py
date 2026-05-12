import io
import json
import os
import re
import sqlite3
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Optional
from urllib.parse import quote

import docx2txt
import numpy as np
import uvicorn
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from openai import AsyncOpenAI
from PIL import Image
from PyPDF2 import PdfReader
from pydantic import BaseModel, Field

# ================= 1. Configuration & Initialization =================
CODE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = CODE_DIR.parent

load_dotenv(PROJECT_ROOT / ".env")
load_dotenv(CODE_DIR / ".env", override=True)

if os.getenv("HF_OFFLINE", "1").lower() in {"1", "true", "yes", "on"}:
    os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
    os.environ.setdefault("HF_DATASETS_OFFLINE", "1")
    os.environ.setdefault("HF_HUB_OFFLINE", "1")


def project_path_from_env(name: str, default: Path) -> Path:
    raw_value = os.getenv(name)
    path = Path(raw_value).expanduser() if raw_value else default
    if not path.is_absolute():
        path = PROJECT_ROOT / path
    return path.resolve()


# Paths
DB_SAVE_PATH = project_path_from_env("CHROMA_DB_PATH", PROJECT_ROOT / "Model" / "chroma_db")
FINANCE_DB_PATH = project_path_from_env("FINANCE_DB_PATH", PROJECT_ROOT / "Model" / "finance_data.db")
LOG_FILE_PATH = project_path_from_env("CHAT_LOG_PATH", PROJECT_ROOT / "Log" / "aegis_chat_logs.jsonl")
EVENT_LOG_PATH = project_path_from_env("EVENT_LOG_PATH", PROJECT_ROOT / "Log" / "platform_events.jsonl")
SERVER_HOST = os.getenv("AEGIS_HOST", "0.0.0.0")
SERVER_PORT = int(os.getenv("AEGIS_PORT", "8000"))

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
if not DEEPSEEK_API_KEY:
    raise ValueError("DEEPSEEK_API_KEY not found. Copy .env.example to Code/.env or project .env first.")

# OCR Engine
ocr_engine = None
ocr_engine_name = None
ocr_init_error = None


def initialize_ocr_engine():
    global ocr_engine, ocr_engine_name, ocr_init_error
    try:
        from rapidocr import RapidOCR
        ocr_engine = RapidOCR()
        ocr_engine_name = "RapidOCR(PP-OCRv4)"
        ocr_init_error = None
        print("[OCR] RapidOCR initialized.")
        return
    except Exception as e:
        ocr_init_error = f"RapidOCR init failed: {e}"
        print(f"[OCR] {ocr_init_error}")
    try:
        import easyocr
        ocr_engine = easyocr.Reader(["ch_sim", "en"])
        ocr_engine_name = "EasyOCR"
        print("[OCR] EasyOCR initialized.")
    except Exception as e:
        ocr_init_error = f"{ocr_init_error}; EasyOCR init failed: {e}" if ocr_init_error else f"EasyOCR init failed: {e}"
        print(f"[OCR] {ocr_init_error}")


initialize_ocr_engine()


def run_ocr(image_np: np.ndarray) -> tuple:
    if ocr_engine is None:
        raise RuntimeError(ocr_init_error or "OCR engine is not initialized")
    if ocr_engine_name and ocr_engine_name.startswith("RapidOCR"):
        result = ocr_engine(image_np)
        lines = [str(line).strip() for line in getattr(result, "txts", ()) if str(line).strip()]
        return "\n".join(lines), None
    else:
        result = ocr_engine.readtext(image_np)
        lines = [item[1] for item in result if item[2] > 0.3]
        avg_conf = float(np.mean([item[2] for item in result])) if result else 0.0
        return "\n".join(lines), avg_conf


# ChromaDB
print("[Chroma] Loading legal vector database...")
try:
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-m3")
    vectordb = Chroma(persist_directory=str(DB_SAVE_PATH), embedding_function=embeddings)
    print(f"[Chroma] Loaded from {DB_SAVE_PATH}")
except Exception as e:
    print(f"[Chroma] Failed to load: {e}")
    vectordb = None

# DeepSeek Client
print("[DeepSeek] Initializing client...")
client = AsyncOpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

# FastAPI App
app = FastAPI(title="Aegis 债优盾", description="中小债权人维权智能平台 — 股东出资义务加速到期专项")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ================= 2. Pydantic Models =================
class ChatRequest(BaseModel):
    query: str
    stream: bool = True
    history: list = []
    top_k: int = 5
    score_threshold: float = 1.5
    scenario: str = "unknown"  # 1-7 或 unknown
    user_name: str = "债权人"


class SourceItem(BaseModel):
    filename: str
    score: float
    content_preview: str


class ChatResponse(BaseModel):
    answer: str
    sources: List[SourceItem] = []


class ExportRequest(BaseModel):
    html_content: str
    filename: str = "债优盾法律文书"


class AdminQuery(BaseModel):
    password: str = "admin888"
    filters: dict = {}


# ================= 3. Utility Functions =================
def save_chat_log(query: str, reasoning: str, answer: str, sources: list, scenario: str = "unknown"):
    log_data = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "user_query": query,
        "justitia_thought": reasoning,
        "justitia_answer": answer,
        "reference_sources": [s['filename'] for s in sources],
        "scenario": scenario,
    }
    try:
        LOG_FILE_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(LOG_FILE_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(log_data, ensure_ascii=False) + "\n")
    except Exception as e:
        print(f"[Log] Write failed: {e}")


def save_platform_event(event_type: str, data: dict):
    try:
        EVENT_LOG_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(EVENT_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps({
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "type": event_type,
                "data": data,
            }, ensure_ascii=False) + "\n")
    except Exception as e:
        print(f"[Event] Write failed: {e}")


# ================= 4. Scenario Definitions =================
SCENARIOS = {
    "1": {
        "title": "公司欠钱不还，想追未实缴股东",
        "description": "公司作为债务人无法清偿到期债务，债权人希望追究未足额缴纳出资的股东责任",
        "legal_basis": ["公司法第54条", "九民纪要第6条", "公司法解释三第13条"],
        "key_elements": ["债权债务关系证明", "公司不能清偿的证明", "股东认缴/实缴出资情况", "股东身份信息"],
    },
    "2": {
        "title": "公司不能清偿，想判断能否主张股东出资加速到期",
        "description": "公司已具备破产原因但不申请破产，债权人想主张股东出资期限加速到期",
        "legal_basis": ["公司法第54条", "九民纪要第6条", "企业破产法第35条"],
        "key_elements": ["债权债务关系证明", "公司不能清偿的证明", "股东出资期限约定", "公司资产负债表/审计报告"],
    },
    "3": {
        "title": "债务发生后，公司延长了股东出资期限",
        "description": "公司对债权人负有债务后，通过股东会决议延长股东出资期限，涉嫌恶意逃避债务",
        "legal_basis": ["公司法第54条", "九民纪要第6条第(2)项"],
        "key_elements": ["债权发生时间的证明", "股东出资期限变更的工商登记", "股东会决议文件"],
    },
    "4": {
        "title": "怀疑股东出资后又把钱转走",
        "description": "股东完成出资验资后，通过虚构交易、关联交易等方式将出资款项转出",
        "legal_basis": ["公司法解释三第12条", "公司法第35条"],
        "key_elements": ["股东出资验资证明", "资金转出银行流水", "关联交易证据", "公司财务账簿"],
    },
    "5": {
        "title": "公司减资后导致债权无法清偿",
        "description": "公司未依法通知债权人即进行减资程序，导致债权人利益受损",
        "legal_basis": ["公司法第177条", "公司法第204条"],
        "key_elements": ["公司减资工商变更记录", "减资公告文件", "未收到减资通知的证明"],
    },
    "6": {
        "title": "股东转让股权后无人履行出资义务",
        "description": "未届出资期限的股东将股权转让给明显无履行能力的主体，逃避出资义务",
        "legal_basis": ["公司法第88条", "公司法解释三第18条"],
        "key_elements": ["股权转让协议/工商变更", "受让方资信状况", "转让时出资期限是否已届满"],
    },
    "7": {
        "title": "不确定属于哪种情况，需要系统初步判断",
        "description": "系统将通过对话引导，帮助债权人梳理案情并自动匹配最合适的维权路径",
        "legal_basis": ["综合适用"],
        "key_elements": ["请尽量描述您遇到的具体情况"],
    },
}


def get_scenario_label(scenario: str) -> str:
    if scenario in SCENARIOS:
        return f"场景{scenario}：{SCENARIOS[scenario]['title']}"
    return "通用法律咨询"


# ================= 5. Intent Router (法财双模) =================
LEGAL_KEYWORDS = [
    '法律', '法条', '条款', '民法', '刑法', '公司法', '劳动法', '合同法', '商法', '合规', '权', '法',
    '章程', '准则', '条例', '规定', '司法解释', '知识产权', '专利', '商标', '著作权',
    '股东', '法人', '董事', '监事', '实控人', '代理人', '原告', '被告', '第三人',
    '连带责任', '有限责任', '股权', '债权', '债务', '担保', '抵押', '质押', '处分',
    '诉讼', '仲裁', '起诉', '申诉', '保全', '判决', '裁定', '调解', '公证', '证据',
    '举证', '质证', '抗辩', '追偿', '执行', '立案', '撤销', '无效', '违约', '侵权',
    '赔偿', '滞纳金', '违约金', '不可抗力', '效力', '判例', '案例', '辩护',
    '协议', '合同', '意向书', '备忘录', '承诺函', '授权书', '通知书', '起诉状', '答辩状',
    '出资', '实缴', '认缴', '加速到期', '减资', '增资', '抽逃', '验资', '出资期限',
    '债权人', '债务人', '欠债', '讨债', '追债', '债权凭证', '不能清偿', '到期债务',
    '股东会', '决议', '扩大出资', '转让股权', '逃避债务',
]

FINANCE_KEYWORDS = [
    '财务', '会计', '报表', '资产', '负债', '权益', '利润', '营收', '成本', '费用',
    '支出', '收入', '科目', '分录', '凭证', '账簿', '对账', '核算', '折旧', '摊销',
    '毛利', '净利', '坏账', '计提', '底稿', '结转', '固定资产', '无形资产', '关联交易',
    '税务', '税收', '纳税', '开票', '发票', '抵扣', '增值税', '所得税', '个税', '企业税',
    '税率', '退税', '印花税', '关税', '核定征收',
    '出资', '实缴', '认缴', '股权', '股票', '融资', '增资', '减资', '股权转让', '期权',
    '收购', '兼并', '清算', '破产', '重整', '对赌', '估值', '尽调', '流水', '套现',
    '盈亏', '赤字', '审计报告', '预算', '决算', '资金', '流动性', '分红',
    '经营', '健康', '状况', '偿债', '能力', '现金流', '毛利率', '净利率', 'ROE', 'ROA',
    '资产负债率', '流动比率', '速动比率', '应收账款', '应付账款',
]


def detect_intent(query: str) -> dict:
    is_legal = any(k in query for k in LEGAL_KEYWORDS)
    is_finance = any(k in query for k in FINANCE_KEYWORDS)
    is_domain = is_legal or is_finance or len(query) > 15
    return {"is_legal": is_legal, "is_finance": is_finance, "is_domain": is_domain}


# ================= 6. Financial DB Query Engine =================
FINANCE_TABLE_SCHEMA = """
本地财务数据库包含以下类型的表：

第一类：标准简易表
- 利润表 (Scode:股票代码, Date:统计日期, REV:营业总收入, NI:净利润, FINEXP:财务费用)
- 资产负债表 (Scode:股票代码, Date:统计日期, CH:货币资金, AT:资产总计, LB:负债合计, EQU:所有者权益合计)
- 现金流量表 (Scode:股票代码, Date:统计日期, NCPOA:经营活动产生的现金流量净额)

第二类：RESSET专业数据库
表名如：RESSET科创板利润表、RESSET新三板资产负债表、RESSET科创板现金流量表、RESSET科创板会计衍生指标 等
RESSET表字段格式为 "中文名称_英文缩写"，例如：
- 基础信息：公司代码_CompanyCode、最新公司全称_LComNm
- 利润表：营业总收入_TotOpRev、净利润_NetProf、归母净利润_NPParentComp
- 资产负债表：资产总计_TotAss、负债合计_TotLiab、货币资金_CashEqv
- 现金流量表：经营活动产生的现金流量净额_NetOpCashFl
- 衍生指标：基本每股收益_BasEPS、每股净资产_NAPS、净资产收益率_ROE

注意：在标准简表中，股票代码查询条件为 Scode = '...'；在 RESSET 表中为 公司代码_CompanyCode = '...'。
"""


async def query_finance_db(query: str) -> dict:
    """AI 自动生成 SQL 并执行，返回财务查账结果"""
    if not FINANCE_DB_PATH.exists():
        return {"sql": None, "result": None, "context": "财务数据库尚未初始化，无本地财务数据可供查询。"}

    sql_prompt = f"""你是精通 SQLite 的数据分析专家。根据以下数据库表结构，为用户的提问写一句 SQL 查询语句。

【表结构与字段映射】
{FINANCE_TABLE_SCHEMA}

【用户提问】：{query}

【严格要求】
1. 只输出一句合法 SQLite 查询，用 ```sql 和 ``` 包裹。不要有任何其他文字。
2. 根据用户提问智能推测表名和列名。
3. 如果无法判断表名或字段，输出 ```sql\nSELECT '无法解析财务查询' AS info;\n```
"""

    try:
        response = await client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": sql_prompt}],
            temperature=0.1,
        )
        sql_text = response.choices[0].message.content
        match = re.search(r"```sql(.*?)```", sql_text, re.DOTALL | re.IGNORECASE)

        if not match:
            return {"sql": None, "result": None, "context": "AI 未能生成有效的 SQL 查询。"}

        sql_query = match.group(1).strip()
        print(f"[Finance SQL] {sql_query}")

        conn = sqlite3.connect(str(FINANCE_DB_PATH))
        cursor = conn.cursor()
        cursor.execute(sql_query)
        rows = cursor.fetchall()
        cols = [desc[0] for desc in cursor.description] if cursor.description else []
        conn.close()

        db_results = [dict(zip(cols, row)) for row in rows]
        context = f"本地财务数据库查到的真实数据（共{len(rows)}条）：{json.dumps(db_results, ensure_ascii=False)}"
        print(f"[Finance] 查询成功，获取 {len(rows)} 条记录")
        return {"sql": sql_query, "result": db_results, "context": context}

    except Exception as e:
        print(f"[Finance] 查询失败: {e}")
        return {"sql": None, "result": None, "context": f"财务数据库查询失败：{str(e)}"}


def get_finance_tables() -> list:
    if not FINANCE_DB_PATH.exists():
        return []
    conn = sqlite3.connect(str(FINANCE_DB_PATH))
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
    tables = [row[0] for row in cursor.fetchall()]
    conn.close()
    return tables


# ================= 7. Case Profile Builder (债权人版本) =================
def build_creditor_case_profile(query: str, scenario: str, user_name: str) -> dict:
    profile = {
        "creditor_name": user_name,
        "debtor_company": "待补充",
        "debt_amount": "待补充",
        "debt_basis": "待补充",
        "shareholder_info": "待补充",
        "subscription_capital": "待补充",
        "paid_capital": "待补充",
        "contribution_deadline": "待补充",
        "company_status": "待补充",
        "existing_evidence": [],
        "scenario": get_scenario_label(scenario),
        "scenario_id": scenario,
    }

    # Simple extraction
    amount_match = re.search(r'(\d+[\.,]?\d*)\s*(万|元|块)', query)
    if amount_match:
        profile["debt_amount"] = amount_match.group(0)

    company_match = re.search(r'(?:公司|被告|债务人)\s*(?:是|为|叫做)?\s*([^\s，。,\.]{2,30}(?:有限公司|股份有限公司|有限责任公司|公司|集团))', query)
    if company_match:
        profile["debtor_company"] = company_match.group(1)

    evidence_types = []
    if any(k in query for k in ['欠条', '借条', '借据']):
        evidence_types.append("债权凭证（欠条/借条/借据）")
    if any(k in query for k in ['合同', '协议']):
        evidence_types.append("合同/协议")
    if any(k in query for k in ['转账', '银行', '流水', '汇款']):
        evidence_types.append("银行转账记录/流水")
    if any(k in query for k in ['微信', '聊天', '短信']):
        evidence_types.append("通讯记录（微信/短信）")
    if any(k in query for k in ['工商', '企查查', '天眼查', '公示']):
        evidence_types.append("工商登记/公示信息")
    if any(k in query for k in ['判决', '裁定', '调解书']):
        evidence_types.append("法院裁判文书")
    if any(k in query for k in ['对账单', '结算单', '确认函']):
        evidence_types.append("对账单/结算单/确认函")
    profile["existing_evidence"] = evidence_types if evidence_types else ["待识别"]

    return profile


# ================= 8. Evidence Gap Analysis =================
def analyze_evidence_gaps(profile: dict, scenario: str) -> dict:
    """根据场景分析证据缺口"""
    required_evidence = {
        "1": ["债权凭证（合同/欠条/判决书）", "公司不能清偿的证明", "工商登记信息（股东认缴/实缴）", "催收记录"],
        "2": ["债权凭证", "公司资产负债表或审计报告", "股东出资期限证明", "公司已具备破产原因的证据"],
        "3": ["债权发生时间证明", "公司延长出资期限的工商变更", "债务发生时的公司章程"],
        "4": ["股东出资验资证明", "资金转出的银行流水", "关联交易合同/凭证"],
        "5": ["公司减资工商变更记录", "减资公告文件", "债权发生时间证明"],
        "6": ["股权转让协议", "工商变更登记", "受让方资信状况证明", "转让时的公司章程"],
        "7": ["请先描述具体案情"],
    }

    needed = required_evidence.get(scenario, required_evidence["1"])
    existing_set = set(profile.get("existing_evidence", []))
    gaps = [e for e in needed if not any(ex in e or e in ex for ex in existing_set)]

    return {
        "required": needed,
        "existing": profile.get("existing_evidence", []),
        "gaps": gaps,
        "completeness": max(0, len(existing_set) / max(len(needed), 1)),
    }


# ================= 9. Core: System Prompt Builder =================
def build_system_prompt(request: ChatRequest, legal_context: str, finance_context: str,
                        case_profile: dict, evidence_analysis: dict) -> str:
    scenario_label = get_scenario_label(request.scenario)
    scenario_info = SCENARIOS.get(request.scenario, SCENARIOS["7"])
    legal_basis_text = "、".join(scenario_info.get("legal_basis", []))
    key_elements_text = "、".join(scenario_info.get("key_elements", []))
    evidence_gaps_text = "、".join(evidence_analysis.get("gaps", [])) or "待用户补充信息后判断"

    return f"""您是"债优盾 Aegis"智能平台的 AI 法律顾问 Justitia，由债优盾团队开发。您的核心使命是帮助中小债权人维护合法权益，特别是通过主张**股东出资义务加速到期**来追讨公司债务。

【系统指令深度对齐】
1. 时间锚点：当前为 2026 年春季。请确保所有法律建议符合最新的 2024 年修订后《公司法》及配套司法解释。
2. 核心法律武器：公司法第54条（股东出资加速到期）、九民纪要第6条、公司法解释三第12-18条、企业破产法第35条。
3. 身份定位：您是专业、冷静、精准的法律顾问，同时富有同理心。服务对象是中小债权人，语言清晰而不晦涩。
4. RAG 驱动：优先引用【本地法律卷宗】中的法条原文、裁判观点和案例。
5. 财务辅助：如果有【本地财务查账结果】，请结合真实数据分析目标公司的偿债能力。

【当前维权场景】
{scenario_label}

【该场景法律依据】
{legal_basis_text}

【该场景关键构成要件】
{key_elements_text}

【结构化案情快照】
{json.dumps(case_profile, ensure_ascii=False, indent=2)}

【证据缺口分析】
需补强证据：{evidence_gaps_text}
证据完整度：{evidence_analysis.get('completeness', 0) * 100:.0f}%

【本地财务查账结果】
{finance_context}

【本地法律卷宗参考】
{legal_context}

【固定输出模板】
除非用户只是问候或简单程序性问题，否则请严格按以下七段式结构输出，不得省略标题：

**一、构成要件审查**
- 根据当前场景，逐一审查股东出资加速到期的法定构成要件是否满足
- 已满足的标记"√"，不明确的标记"待查"
- 引用公司法第54条及九民纪要的对应条款

**二、偿债能力初步测算**
- 如果查询到财务数据，用具体数字分析目标公司偿债能力
- 如果无财务数据，说明哪些指标需要用户进一步提供
- 给出偿债能力等级：健康 / 一般 / 严重不足 / 资不抵债

**三、证据链研判**
- 先给结论：证据较充分 / 已有初步证据但需补强 / 证据严重不足
- 逐一列出已有证据和缺失证据
- 为每个缺失证据提供替代性取证方案

**四、加速到期可行性评分**
- 综合评估后给出 0-100 分的可行性评分
- 说明评分依据和扣分项
- 如果评分偏低，说明通过补哪些证据可以提升

**五、可走的维权路径**
- 按层次说明：协商/发函 → 诉前财产保全 → 诉讼（股东出资加速到期之诉）→ 执行
- 说明每种路径的优缺点和时效要求
- 提示是否需要申请财产保全或证据保全

**六、下一步请您先做这几件事**
- 给 3 到 5 个可立即执行的动作，按优先级排列
- 每个动作说明为什么重要

**七、还需要您补充的信息**
- 只问最影响案件推进的 3 到 5 个问题
- 每个问题解释为什么需要这个信息

【交互准则】
- 严禁提及您的 AI 架构、训练截止日期或底层模型
- 对数字和金额务必精确，不确定时标注"待核实"
- 如果用户上传的 OCR 结果模糊，请委婉请其通过文字补充关键数字
- 既保持法律专业度，又让普通债权人能听懂
- 始终以中文回答，排版利于电脑端和手机端阅读
- 当前用户：{request.user_name}

Respond strictly in Chinese.
"""


# ================= 10. DOCX Export (债权人版本) =================
def clean_html_summary(html: str) -> str:
    if not html:
        return "债优盾法律文书"
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "svg", "img", "button", "input"]):
        tag.decompose()
    for tag in soup.find_all(True):
        for attr in list(tag.attrs):
            if attr not in ("colspan", "rowspan"):
                del tag[attr]
    cleaned = str(soup)
    cleaned = re.sub(r'\n\s*\n', '\n', cleaned)
    return (cleaned or "债优盾法律文书")[:80]


def export_creditor_complaint_docx(html_content: str, filename: str = "债优盾法律文书") -> io.BytesIO:
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("民 事 起 诉 状")
    run.font.size = Pt(16)
    run.bold = True

    # Parse from markdown
    text_parts = []
    for el in soup.descendants:
        if isinstance(el, NavigableString):
            t = str(el).strip()
            if t and t not in text_parts:
                text_parts.append(t)

    full_text = "\n".join(text_parts)
    paragraphs = [p.strip() for p in full_text.split("\n") if len(p.strip()) > 5]

    for p_text in paragraphs[:50]:
        para = doc.add_paragraph()
        if "原告" in p_text or "申请人" in p_text:
            para.paragraph_format.first_line_indent = Cm(0)
        else:
            para.paragraph_format.first_line_indent = Cm(0.74)
        run = para.add_run(p_text)
        run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)[:80]
    return buf


# ================= 11. API Endpoints =================
@app.get("/api/health")
async def health_check():
    tables = get_finance_tables()
    return {
        "status": "ok",
        "service": "Aegis 债优盾",
        "ocr_engine": ocr_engine_name,
        "chroma_loaded": vectordb is not None,
        "finance_tables": len(tables),
        "scenarios": len(SCENARIOS),
    }


@app.post("/api/chat")
async def chat_endpoint(request: ChatRequest):
    intent = detect_intent(request.query)
    finance_context = "用户问题未涉及具体财务指标查询，或本地财务库无相关记录。"
    legal_context = "用户问题未涉及具体法律案卷，无需引用本地判例。"
    source_items = []
    sql_meta = None

    # Build case profile
    case_profile = build_creditor_case_profile(request.query, request.scenario, request.user_name)
    evidence_analysis = analyze_evidence_gaps(case_profile, request.scenario)

    # Intent routing
    if intent["is_domain"]:
        print(f"[Intent] Legal={intent['is_legal']}, Finance={intent['is_finance']}")

        # Financial query
        if intent["is_finance"]:
            print("[Finance] Triggering SQL generation...")
            finance_result = await query_finance_db(request.query)
            finance_context = finance_result["context"]
            sql_meta = {"sql": finance_result["sql"], "result": finance_result["result"]}

        # Legal RAG
        if intent["is_legal"] and vectordb is not None:
            print("[Legal] Searching ChromaDB...")
            try:
                raw_results = vectordb.similarity_search_with_score(request.query, k=request.top_k)
                legal_context = ""
                for i, (doc, score) in enumerate(raw_results):
                    if score < request.score_threshold:
                        filename = doc.metadata.get('source', '未知')
                        category = doc.metadata.get('category_path', '')
                        legal_context += f"\n--- 案卷{i + 1} (来源: {category}/{filename}) ---\n{doc.page_content}\n"
                        source_items.append({
                            "filename": f"{category}/{filename}" if category else filename,
                            "score": round(score, 4),
                            "content_preview": doc.page_content[:50] + "..."
                        })
            except Exception as e:
                print(f"[Legal] RAG failed: {e}")
                legal_context = f"法律知识库检索异常：{str(e)}"

    # Build system prompt
    final_system_prompt = build_system_prompt(
        request, legal_context, finance_context, case_profile, evidence_analysis
    )

    messages = [{"role": "system", "content": final_system_prompt}]
    if request.history:
        for msg in request.history:
            if isinstance(msg, dict) and "role" in msg and "content" in msg:
                messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": request.query})

    if not request.stream:
        # Non-streaming
        response = await client.chat.completions.create(
            model="deepseek-chat",
            messages=messages,
            max_tokens=8192,
        )
        answer = response.choices[0].message.content
        save_chat_log(request.query, "", answer, source_items, request.scenario)
        return {"answer": answer, "sources": source_items, "case_profile": case_profile}

    # Streaming
    async def generate_stream():
        meta = {"type": "meta", "sources": source_items, "case_profile": case_profile}
        if sql_meta:
            meta["finance"] = sql_meta
        yield f"data: {json.dumps(meta, ensure_ascii=False)}\n\n"

        accumulated_reasoning = ""
        accumulated_content = ""

        try:
            response = await client.chat.completions.create(
                model="deepseek-chat",
                messages=messages,
                stream=True,
                max_tokens=8192,
            )
            async for chunk in response:
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta
                if hasattr(delta, 'reasoning_content') and delta.reasoning_content:
                    accumulated_reasoning += delta.reasoning_content
                    yield f"data: {json.dumps({'type': 'reasoning', 'content': delta.reasoning_content}, ensure_ascii=False)}\n\n"
                if hasattr(delta, 'content') and delta.content:
                    accumulated_content += delta.content
                    yield f"data: {json.dumps({'type': 'chunk', 'content': delta.content}, ensure_ascii=False)}\n\n"
        except Exception as e:
            print(f"[LLM Error] {e}")
            fallback = f"抱歉，AI 服务暂时不可用（{str(e)[:100]}）。请稍后重试。"
            accumulated_content += fallback
            yield f"data: {json.dumps({'type': 'chunk', 'content': fallback}, ensure_ascii=False)}\n\n"

        save_chat_log(request.query, accumulated_reasoning, accumulated_content, source_items, request.scenario)
        yield "data: [DONE]\n\n"

    return StreamingResponse(generate_stream(), media_type="text/event-stream")


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    content = await file.read()
    ext = Path(file.filename).suffix.lower()

    extracted_text = ""
    ocr_confidence = None

    if ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff"):
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(content)).convert("RGB")
            img_np = np.array(img)
            extracted_text, ocr_confidence = run_ocr(img_np)
        except Exception as e:
            raise HTTPException(500, f"OCR 处理失败：{str(e)}")

    elif ext == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            t = page.extract_text()
            if t:
                extracted_text += t + "\n"

    elif ext in (".docx", ".doc"):
        extracted_text = docx2txt.process(io.BytesIO(content))

    elif ext == ".txt":
        try:
            extracted_text = content.decode("utf-8")
        except UnicodeDecodeError:
            extracted_text = content.decode("gbk")

    else:
        raise HTTPException(400, f"不支持的文件格式：{ext}")

    extracted_text = extracted_text.strip()[:8000]
    save_platform_event("file_upload", {
        "filename": file.filename,
        "ext": ext,
        "text_length": len(extracted_text),
        "ocr_confidence": ocr_confidence,
    })

    return {
        "filename": file.filename,
        "text": extracted_text,
        "confidence": ocr_confidence,
        "engine": ocr_engine_name,
    }


@app.post("/api/export-docx")
async def export_docx(request: ExportRequest):
    buf = export_creditor_complaint_docx(request.html_content, request.filename)
    safe_filename = re.sub(r'[<>:"/\\|?*]', '_', request.filename)[:80]
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe_filename)}.docx"}
    )


@app.get("/api/finance/tables")
async def list_finance_tables():
    tables = get_finance_tables()
    db_exists = FINANCE_DB_PATH.exists()
    return {"db_exists": db_exists, "tables": tables, "count": len(tables)}


@app.get("/api/scenarios")
async def get_scenarios():
    return SCENARIOS


# ================= 12. Admin Dashboard =================
@app.get("/api/admin/events")
async def get_admin_events(password: str = ""):
    events = []
    if EVENT_LOG_PATH.exists():
        with open(EVENT_LOG_PATH, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    events.append(json.loads(line))
    return {"events": events[-200:], "total": len(events)}


# ================= 13. Static Frontend =================
WEB_DIR = CODE_DIR / "Web"


@app.get("/")
async def serve_frontend():
    if (WEB_DIR / "index.html").exists():
        return FileResponse(WEB_DIR / "index.html")
    raise HTTPException(404, "Frontend not found")


@app.get("/{path:path}")
async def serve_static(path: str):
    file_path = WEB_DIR / path
    if file_path.exists() and file_path.is_file():
        return FileResponse(file_path)
    # Fallback to SPA
    if (WEB_DIR / "index.html").exists():
        return FileResponse(WEB_DIR / "index.html")
    raise HTTPException(404, "Not found")


# ================= 14. Entry Point =================
if __name__ == "__main__":
    print(f"\n{'='*60}")
    print(f"  Aegis 债优盾 — 中小债权人维权智能平台")
    print(f"  股东出资义务加速到期专项 AI 法律顾问")
    print(f"{'='*60}\n")
    uvicorn.run(app, host=SERVER_HOST, port=SERVER_PORT, log_level="info")
